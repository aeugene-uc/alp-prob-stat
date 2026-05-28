# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- base style: Times New Roman 10pt ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = normal.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 1.0

# page margins (JUISI-ish single column)
for s in doc.sections:
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

def _set_run(r, size=10, bold=False, italic=False):
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    rpr = r._element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn('w:eastAsia'), 'Times New Roman')

def para(text='', align='justify', size=10, bold=False, italic=False,
         indent=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    a = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
         'left': WD_ALIGN_PARAGRAPH.LEFT, 'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.alignment = a
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    if text:
        r = p.add_run(text); _set_run(r, size, bold, italic)
    return p

def runs_para(segments, align='justify', indent=False, space_after=6):
    """segments: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    for t, b, i in segments:
        r = p.add_run(t); _set_run(r, 10, b, i)
    return p

def heading(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num}. {text}' if num else text)
    _set_run(r, 10, bold=True)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); _set_run(r, 10, bold=True, italic=True)
    return p

def _horizontal_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom', 'insideH'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ('left', 'right', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)

def table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); _set_run(r, 10, bold=False)

def make_table(rows, header=True, widths=None, align_center=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _horizontal_borders(t)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(val))
            _set_run(r, 9, bold=(header and ri == 0))
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(rows)):
                t.cell(ri, ci).width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, caption, width=5.8):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(); r.add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    rr = c.add_run(caption); _set_run(rr, 9, bold=False)

# =====================================================================
# TITLE BLOCK
# =====================================================================
para('Pemodelan Konsumsi Bahan Bakar Kendaraan Menggunakan Regresi '
     'Polinomial Bertingkat dengan Seleksi Fitur Berbasis Korelasi',
     align='center', size=16, bold=True, space_after=4)
para('Modeling Vehicle Fuel Consumption Using Sequential Polynomial '
     'Regression with Correlation-Based Feature Selection',
     align='center', size=12, bold=False, italic=True, space_after=8)

runs_para([('Amadeus Eugene Dirgantara', False, False), ('1', False, False),
           (', Rei Putra Soemanto', False, False), ('2', False, False),
           (', Jason Tio', False, False), ('3', False, False),
           (', Christopher Andreas, S.Stat., M.Stat.', False, False), ('4', False, False)],
          align='center', space_after=2)
para('1,2,3,4 Program Studi Informatika, Universitas Ciputra, Surabaya, '
     'Jawa Timur, Indonesia', align='center', size=9, space_after=2)
para('*Email penulis korespondensi: christopher.andreas@ciputra.ac.id',
     align='center', size=9, space_after=8)

# ----- Abstrak (Indonesia) -----
runs_para([
    ('Abstrak: ', True, False),
    ('Pemilihan bentuk fungsional yang tepat merupakan tahap krusial dalam '
     'analisis regresi. Regresi linear sering diterapkan sebagai pilihan default, '
     'padahal intuisi rekayasa menunjukkan bahwa hubungan antara ukuran mesin dan '
     'konsumsi bahan bakar bersifat nonlinier akibat menurunnya efisiensi pembakaran '
     'pada kapasitas mesin yang lebih besar. Penelitian ini memodelkan konsumsi bahan '
     'bakar gabungan kendaraan (L/100 km) menggunakan dataset Fuel Consumption 2000–2022 '
     '(22.555 baris setelah pembersihan). Kerangka kerja keputusan bertingkat diterapkan: '
     'Tahap 1 melakukan seleksi fitur berbasis korelasi Pearson dan Spearman, Tahap 2 dan 3 '
     'menjalankan Nested F-Test untuk menentukan orde polinomial terendah yang memadai, dan '
     'Tahap 4 memvalidasi asumsi Ordinary Least Squares (OLS). Seleksi fitur menyisakan satu '
     'prediktor, yaitu ukuran mesin, karena jumlah silinder hampir kolinier (r ≈ 0,91) dan '
     'tahun model nyaris tidak berkorelasi dengan respon. Kedua Nested F-Test menolak hipotesis '
     'nol secara meyakinkan sehingga model kubik terpilih sebagai model final. Model menjelaskan '
     'sekitar 66% variansi konsumsi bahan bakar (R² uji ≈ 0,658) dengan RMSE ≈ 1,73 L/100 km '
     'dan tanpa indikasi overfitting. Uji normalitas dan homoskedastisitas ditolak secara formal, '
     'namun hal ini merupakan konsekuensi wajar dari ukuran sampel besar dan tidak membatalkan '
     'estimasi titik koefisien. Temuan ini menegaskan adanya pola diminishing returns pada '
     'hubungan ukuran mesin dengan konsumsi bahan bakar.', False, False)
], space_after=4)
runs_para([('Kata Kunci: ', True, False),
           ('regresi polinomial, Nested F-Test, seleksi fitur, konsumsi bahan bakar, '
            'multikolinieritas, validasi asumsi OLS.', False, False)], space_after=6)

runs_para([
    ('Abstract: ', True, True),
    ('Selecting the correct functional form is a crucial step in regression analysis. '
     'Linear regression is often applied by default, even though engineering intuition '
     'suggests that the relationship between engine size and fuel consumption is nonlinear '
     'due to declining combustion efficiency at larger displacements. This study models the '
     'combined fuel consumption of vehicles (L/100 km) using the Fuel Consumption 2000–2022 '
     'dataset (22,555 rows after cleaning). A sequential decision framework is applied: Stage 1 '
     'performs correlation-based feature selection using Pearson and Spearman coefficients, '
     'Stages 2 and 3 run Nested F-Tests to determine the lowest adequate polynomial order, and '
     'Stage 4 validates the Ordinary Least Squares (OLS) assumptions. Feature selection retains a '
     'single predictor, engine size, because cylinder count is nearly collinear (r ≈ 0.91) and '
     'model year is virtually uncorrelated with the response. Both Nested F-Tests reject the null '
     'hypothesis convincingly, so the cubic model is selected as the final model. The model explains '
     'about 66% of fuel-consumption variance (test R² ≈ 0.658) with an RMSE ≈ 1.73 L/100 km and no '
     'sign of overfitting. Normality and homoscedasticity tests are formally rejected, but this is an '
     'expected consequence of the large sample size and does not invalidate the coefficient point '
     'estimates. The findings confirm a diminishing-returns pattern in the engine-size–fuel-consumption '
     'relationship.', False, True)
], space_after=4)
runs_para([('Keywords: ', True, True),
           ('polynomial regression, Nested F-Test, feature selection, fuel consumption, '
            'multicollinearity, OLS assumption validation.', False, True)], space_after=8)

# =====================================================================
# 1. PENDAHULUAN
# =====================================================================
heading(1, 'Pendahuluan')
para('Konsumsi bahan bakar kendaraan merupakan indikator penting yang berkaitan langsung '
     'dengan biaya operasional, emisi gas rumah kaca, dan kebijakan keberlanjutan transportasi. '
     'Kemampuan memprediksi konsumsi bahan bakar dari karakteristik teknis kendaraan memiliki '
     'nilai praktis bagi perencanaan armada, perbandingan biaya konsumen, deteksi anomali pada '
     'data spesifikasi, hingga pelaporan keberlanjutan. Oleh karena itu, pemodelan statistik yang '
     'tepat terhadap hubungan antara karakteristik kendaraan dan konsumsi bahan bakar menjadi '
     'kebutuhan yang relevan.', indent=True)
para('Dalam praktik, regresi linear kerap dipilih sebagai model default karena kesederhanaan '
     'dan kemudahan interpretasinya. Namun, intuisi fisik menyatakan bahwa hubungan antara ukuran '
     'mesin dan konsumsi bahan bakar tidak selalu linear. Pada kapasitas mesin yang semakin besar, '
     'efisiensi pembakaran cenderung menurun sehingga tambahan konsumsi per satuan kenaikan ukuran '
     'mesin mengecil. Pola ini dikenal sebagai diminishing returns dan secara matematis menghasilkan '
     'kurva yang melengkung (konkaf). Model yang dispesifikasikan secara keliru berisiko mengalami '
     'underfitting, yaitu melewatkan kelengkungan nyata dalam data, atau overfitting, yaitu '
     'memperkenalkan suku berorde tinggi yang sebenarnya palsu.', indent=True)
para('Tantangan kedua adalah pemilihan prediktor. Beberapa karakteristik kendaraan saling '
     'berkorelasi tinggi, misalnya ukuran mesin dan jumlah silinder. Memasukkan prediktor yang hampir '
     'kolinier ke dalam satu model akan menyebabkan multikolinieritas, yang membuat standard error '
     'membengkak dan estimasi koefisien menjadi tidak stabil. Dengan demikian, seleksi fitur yang '
     'disiplin perlu dilakukan sebelum pemodelan.', indent=True)
runs_para([
    ('Penelitian ini bertujuan menjawab tiga pertanyaan penelitian yang saling berkaitan. ', False, False),
    ('Pertama', True, False),
    (', fitur kendaraan kontinu mana yang sebaiknya dipertahankan sebagai prediktor. ', False, False),
    ('Kedua', True, False),
    (', berapakah polinomial berorde terendah yang secara memadai memodelkan hubungan tersebut '
     'dengan konsumsi bahan bakar gabungan. ', False, False),
    ('Ketiga', True, False),
    (', apakah model yang dihasilkan memenuhi asumsi regresi Ordinary Least Squares (OLS). '
     'Untuk menjawabnya, penelitian ini menerapkan kerangka keputusan bertingkat yang menggabungkan '
     'seleksi fitur berbasis korelasi, pengujian hipotesis bersarang (Nested F-Test) untuk menentukan '
     'orde polinomial, serta validasi asumsi secara formal dan visual.', False, False)
], indent=True)

# =====================================================================
# 2. KAJIAN PUSTAKA
# =====================================================================
heading(2, 'Kajian Pustaka')
subheading('2.1 Regresi Linear dan Polinomial')
para('Regresi linear berganda memodelkan respon y sebagai kombinasi linear dari prediktor. '
     'Ketika hubungan antara prediktor dan respon melengkung namun tetap monotonik, regresi '
     'polinomial menjadi perluasan yang wajar. Regresi polinomial tetap linear pada parameternya '
     'sehingga dapat diestimasi dengan OLS, namun mampu menangkap kelengkungan melalui penambahan '
     'suku berpangkat, seperti x² dan x³. Model polinomial orde-p untuk satu prediktor dituliskan '
     'sebagai y = β₀ + β₁x + β₂x² + … + βₚxᵖ + ε.', indent=True)
subheading('2.2 Korelasi Pearson dan Spearman')
para('Koefisien korelasi Pearson mengukur kekuatan asosiasi linear antara dua variabel, sedangkan '
     'koefisien korelasi Spearman mengukur asosiasi monotonik tanpa mengasumsikan bentuk hubungan. '
     'Perbandingan keduanya bersifat informatif: apabila Spearman secara nyata lebih tinggi daripada '
     'Pearson pada pasangan prediktor–respon, kesenjangan tersebut menjadi indikasi bahwa hubungannya '
     'monotonik tetapi tidak linear. Kondisi inilah yang tepat ditangani oleh perluasan polinomial.', indent=True)
subheading('2.3 Multikolinieritas')
para('Multikolinieritas terjadi ketika dua atau lebih prediktor saling berkorelasi sangat tinggi. '
     'Konsekuensinya, matriks desain mendekati singular sehingga estimasi koefisien menjadi sensitif '
     'terhadap perubahan kecil pada data dan standard error membesar. Salah satu strategi penanganannya '
     'adalah mempertahankan prediktor yang lebih kuat berkaitan dengan respon dan membuang pasangannya '
     'yang lebih lemah.', indent=True)
subheading('2.4 Nested F-Test dan Validasi Asumsi OLS')
para('Nested F-Test membandingkan dua model bersarang, yaitu model yang lebih sederhana dan '
     'perluasannya yang menambah satu atau beberapa suku. Uji ini menilai apakah pengurangan jumlah '
     'kuadrat residu (Residual Sum of Squares, RSS) akibat penambahan suku cukup besar untuk '
     'membenarkan parameter tambahan. Inferensi OLS bertumpu pada beberapa asumsi, terutama '
     'normalitas residu dan homoskedastisitas (ragam residu konstan). Normalitas dapat diuji dengan '
     'Anderson-Darling test, sedangkan homoskedastisitas dengan Breusch-Pagan test, keduanya dilengkapi '
     'diagnostik visual berupa Q-Q plot dan plot residuals vs fitted.', indent=True)

# =====================================================================
# 3. METODE
# =====================================================================
heading(3, 'Metode')
subheading('3.1 Dataset dan Variabel')
para('Penelitian ini menggunakan dataset publik Fuel Consumption 2000–2022 yang memuat 22.556 '
     'baris dan 13 kolom spesifikasi kendaraan. Variabel respon adalah konsumsi bahan bakar gabungan '
     'comb_l_100km (diambil dari kolom COMB (L/100 km)). Tiga kandidat prediktor kontinu '
     'dipertimbangkan, yaitu ukuran mesin (engine_size, dalam liter), jumlah silinder (cylinders), dan '
     'tahun model (year).', indent=True)
para('Sebuah iterasi awal sempat mengekstrak jumlah gigi dari kode transmisi sebagai kandidat '
     'prediktor keempat. Pendekatan tersebut ditinggalkan karena dua alasan. Pertama, korelasi '
     'marginalnya dengan respon mendekati nol. Kedua, ekstraksi memaksa pembuangan baris kendaraan '
     'bertransmisi CVT, sehingga membiaskan sampel. Mempertahankan seluruh baris dan menghapus fitur '
     'tersebut menghasilkan dataset yang tidak bias terhadap tipe transmisi.', indent=True)
subheading('3.2 Persiapan Data')
para('Data mentah diperiksa untuk nilai yang hilang dan tidak ditemukan satu pun nilai kosong pada '
     'seluruh kolom, sehingga tidak diperlukan imputasi maupun penghapusan baris berbasis missing '
     'value. Selanjutnya, duplikat persis dihapus karena produsen kerap mengirimkan banyak konfigurasi '
     'trim yang identik dan dataset mencatat tiap trim sebagai baris tersendiri. Sebanyak satu baris '
     'duplikat dihapus sehingga tersisa 22.555 baris sebagai ukuran sampel efektif.', indent=True)
subheading('3.3 Kerangka Keputusan Bertingkat')
para('Analisis dilakukan dalam empat tahap berurutan:', indent=False, space_after=3)
for txt in [
    ('Tahap 1 — Seleksi fitur.', ' Buang prediktor yang korelasi marginalnya lemah terhadap respon '
     '(|ρ| < 0,20 pada Pearson maupun Spearman) atau yang hampir kolinier dengan prediktor lain yang '
     'lebih kuat (|r| > 0,85).'),
    ('Tahap 2 — Uji suku kuadratik.', ' Nested F-Test membandingkan model linear (M₁) dengan model '
     'kuadratik (M₂). H₀: β₂ = 0.'),
    ('Tahap 3 — Uji suku kubik.', ' Dijalankan hanya bila H₀ Tahap 2 ditolak. Nested F-Test '
     'membandingkan M₂ dengan model kubik (M₃). H₀: β₃ = 0.'),
    ('Tahap 4 — Validasi asumsi.', ' Anderson-Darling test untuk normalitas residu dan Breusch-Pagan '
     'test untuk homoskedastisitas, dilengkapi Q-Q plot serta plot residuals vs fitted.'),
]:
    runs_para([(txt[0], True, False), (txt[1], False, False)], indent=True, space_after=3)
para('Kandidat model didefinisikan sebagai berikut, dengan x menyatakan ukuran mesin:', space_after=3)
para('M₁ (linear):    y = β₀ + β₁x + ε', align='center', space_after=2)
para('M₂ (kuadratik): y = β₀ + β₁x + β₂x² + ε', align='center', space_after=2)
para('M₃ (kubik):     y = β₀ + β₁x + β₂x² + β₃x³ + ε', align='center', space_after=6)
subheading('3.4 Estimasi dan Evaluasi')
para('Data dibagi menjadi subset pelatihan dan uji dengan rasio 80/20 (18.044 baris pelatihan dan '
     '4.511 baris uji). Subset pelatihan digunakan untuk fitting model dan Nested F-Test, sedangkan '
     'subset uji ditahan untuk menilai generalisasi. Ketiga model diestimasi dengan OLS. Kualitas model '
     'dibandingkan melalui R², Adjusted R², AIC, BIC, dan RSS, sedangkan akurasi generalisasi diukur '
     'dengan R², RMSE, dan MAE pada subset uji. Seluruh komputasi dilakukan dengan pustaka Python '
     '(pandas, statsmodels, scikit-learn) menggunakan seed acak tetap demi reproduktibilitas.', indent=True)

# =====================================================================
# 4. HASIL DAN PEMBAHASAN
# =====================================================================
heading(4, 'Hasil dan Pembahasan')

subheading('4.1 Statistik Deskriptif')
para('Tabel 1 menyajikan statistik deskriptif variabel respon dan kandidat prediktor. Ukuran mesin '
     'berkisar 0,8–8,4 L dengan median 3,0 L, jumlah silinder terkonsentrasi pada konfigurasi 4/6/8, '
     'dan tahun model tersebar cukup merata pada rentang 2000–2022. Respon comb_l_100km memiliki '
     'skewness kanan dengan mayoritas kendaraan pada rentang 8–14 L/100 km dan ekor tipis melampaui '
     '20 L/100 km.', indent=True)
table_caption('Tabel 1. Statistik deskriptif variabel (n = 22.555)')
make_table([
    ['Variabel', 'Mean', 'Std', 'Min', 'Q1', 'Median', 'Q3', 'Max'],
    ['engine_size (L)', '3,357', '1,335', '0,8', '2,3', '3,0', '4,2', '8,4'],
    ['cylinders', '5,854', '1,820', '2', '4', '6', '8', '16'],
    ['year', '2011,6', '6,298', '2000', '2006', '2012', '2017', '2022'],
    ['comb_l_100km', '11,034', '2,911', '3,6', '9,1', '10,6', '12,7', '26,1'],
], widths=[1.5, 0.7, 0.7, 0.6, 0.6, 0.75, 0.6, 0.6])

subheading('4.2 Distribusi dan Pola Visual')
para('Gambar 1 memperlihatkan distribusi keempat variabel. Ukuran mesin dan konsumsi bahan bakar '
     'sama-sama memiliki skewness kanan dengan ekor atas panjang yang berasal dari kendaraan mewah dan '
     'berperforma tinggi. Jumlah silinder bersifat diskret dan terkonsentrasi pada 4, 6, dan 8 silinder, '
     'sedangkan tahun model terdistribusi hampir uniform.', indent=True)
figure('figures/cell18_img1.png', 'Gambar 1. Distribusi ukuran mesin, jumlah silinder, tahun model, '
       'dan konsumsi bahan bakar gabungan.', width=5.6)
para('Gambar 2 menampilkan plot pencar setiap prediktor terhadap respon. Ukuran mesin menunjukkan '
     'hubungan monotonik yang jelas namun melandai setelah sekitar 5 L, yaitu bentuk kelengkungan '
     'konkaf yang konsisten dengan pola diminishing returns. Jumlah silinder juga sangat terkait dengan '
     'respon namun bersifat diskret sehingga membentuk garis-garis vertikal. Tahun model tampak datar, '
     'mengindikasikan korelasi yang nyaris nol dengan respon.', indent=True)
figure('figures/cell21_img2.png', 'Gambar 2. Plot pencar tiap prediktor terhadap konsumsi bahan bakar '
       'gabungan. Kelengkungan pada ukuran mesin menjadi motivasi visual perluasan polinomial.', width=6.2)

subheading('4.3 Analisis Korelasi dan Seleksi Fitur (Tahap 1)')
para('Tabel 2 merangkum korelasi Pearson dan Spearman setiap prediktor terhadap respon, beserta '
     'kesenjangan absolutnya. Matriks korelasi lengkap disajikan pada Lampiran (Gambar A1 dan A2). Dua '
     'prediktor menunjukkan asosiasi kuat dengan respon, yaitu ukuran mesin (r ≈ 0,807) dan jumlah '
     'silinder (r ≈ 0,772), sedangkan tahun model praktis tidak berkorelasi (|r| < 0,07). Kesenjangan '
     'positif yang berarti pada ukuran mesin (Spearman − Pearson ≈ 0,038) menguatkan indikasi '
     'kelengkungan yang terlihat pada plot pencar.', indent=True)
table_caption('Tabel 2. Korelasi prediktor terhadap respon comb_l_100km')
make_table([
    ['Prediktor', 'Pearson', 'Spearman', '|Spearman| − |Pearson|'],
    ['engine_size', '0,807', '0,845', '0,038'],
    ['cylinders', '0,772', '0,818', '0,046'],
    ['year', '−0,068', '−0,073', '0,005'],
], widths=[1.6, 1.0, 1.0, 1.8])
para('Hal yang patut diwaspadai adalah korelasi antara ukuran mesin dan jumlah silinder yang sangat '
     'tinggi (Pearson r ≈ 0,91; Spearman ≈ 0,94). Nilai sebesar itu menandakan kedua prediktor hampir '
     'merupakan variabel yang sama, sehingga memasukkan keduanya ke dalam OLS akan menimbulkan '
     'multikolinieritas. Aturan Tahap 1 menanganinya dengan mempertahankan prediktor yang lebih kuat '
     '(ukuran mesin) dan membuang pasangannya (jumlah silinder). Tabel 3 mendokumentasikan keputusan '
     'seleksi fitur secara teraudit.', indent=True)
table_caption('Tabel 3. Keputusan seleksi fitur Tahap 1')
make_table([
    ['Prediktor', 'r dengan y', 'Peer terkuat', 'r dengan peer', 'Keputusan'],
    ['engine_size', '0,807', 'cylinders', '0,913', 'PERTAHANKAN (lebih kuat ke y)'],
    ['cylinders', '0,772', 'engine_size', '0,913', 'BUANG (kolinier r≈0,91)'],
    ['year', '−0,068', 'engine_size', '−0,078', 'BUANG (korelasi lemah)'],
], widths=[1.2, 0.9, 1.1, 1.0, 2.1], align_center=True)
para('Setelah Tahap 1, hanya ukuran mesin yang dipertahankan. Akibatnya, analisis menyusut menjadi '
     'regresi polinomial dengan satu prediktor dan ketiga kandidat model menyederhana sesuai definisi '
     'pada Subbab 3.3.', indent=True)

subheading('4.4 Fitting Model dan Pemilihan Orde Polinomial (Tahap 2–3)')
para('Tabel 4 membandingkan ketiga kandidat model pada subset pelatihan. R² naik secara monoton seiring '
     'penambahan suku, sebagaimana diharapkan, namun kenaikan dari M₂ ke M₃ jauh lebih kecil daripada '
     'dari M₁ ke M₂. AIC dan BIC sama-sama terendah pada M₃, mengisyaratkan bahwa kompleksitas tambahan '
     'masih terbayar.', indent=True)
table_caption('Tabel 4. Perbandingan kandidat model (data pelatihan, n = 18.044)')
make_table([
    ['Model', 'k', 'R²', 'Adj R²', 'AIC', 'BIC', 'RSS'],
    ['M₁ linear', '2', '0,6511', '0,6511', '70632,6', '70648,2', '52941,1'],
    ['M₂ kuadratik', '3', '0,6565', '0,6565', '70353,8', '70377,2', '52123,4'],
    ['M₃ kubik', '4', '0,6570', '0,6569', '70329,9', '70361,1', '52048,9'],
], widths=[1.3, 0.4, 0.8, 0.8, 1.0, 1.0, 1.0])
para('Keputusan formal diambil melalui dua Nested F-Test yang dirangkum pada Tabel 5. Pada Tahap 2, '
     'penambahan suku kuadratik menghasilkan F = 283,0 dengan p ≈ 5,0 × 10⁻⁶³, sehingga H₀ ditolak dan '
     'spesifikasi kuadratik lebih disukai daripada linear. Karena Tahap 2 menolak H₀, Tahap 3 dijalankan. '
     'Penambahan suku kubik menghasilkan F = 25,8 dengan p ≈ 3,7 × 10⁻⁷, sehingga H₀ kembali ditolak. '
     'Dengan demikian, model kubik (M₃) terpilih sebagai model final.', indent=True)
table_caption('Tabel 5. Ringkasan Nested F-Test (Tahap 2 dan 3)')
make_table([
    ['Perbandingan', 'ΔRSS', 'df', 'F', 'p-value', 'Keputusan'],
    ['M₁ → M₂ (Tahap 2)', '817,65', '1', '283,0', '5,0 × 10⁻⁶³', 'Tolak H₀'],
    ['M₂ → M₃ (Tahap 3)', '74,56', '1', '25,8', '3,7 × 10⁻⁷', 'Tolak H₀'],
], widths=[1.7, 0.9, 0.4, 0.7, 1.2, 1.1])
para('Gambar 3 menumpangkan ketiga kurva polinomial pada data pelatihan. Kurva kubik terpilih (garis '
     'tebal) mengikuti pusat awan titik di sepanjang rentang ukuran mesin: naik tajam pada mesin kecil '
     'lalu melandai pada mesin besar. Perbedaan yang jelas antara kurva linear dan kurva orde lebih '
     'tinggi mengonfirmasi bahwa kelengkungan memang penting, konsisten dengan hasil Nested F-Test.', indent=True)
figure('figures/cell60_img6.png', 'Gambar 3. Ukuran mesin terhadap konsumsi bahan bakar gabungan dengan '
       'kurva polinomial M₁, M₂, dan M₃. Kurva kubik terpilih digambar tebal.', width=5.4)

subheading('4.5 Model Final dan Interpretasi Koefisien')
para('Tabel 6 menyajikan koefisien model kubik final beserta standard error, statistik t, p-value, dan '
     'selang kepercayaan 95%. Seluruh koefisien polinomial berbeda nyata dari nol pada α = 0,05. Tanda '
     'koefisien menghasilkan kurva yang naik lalu melandai pada rentang ukuran mesin teramati, yaitu '
     'wujud pola diminishing returns yang diprediksi secara fisik. Intercept tidak diinterpretasikan '
     'langsung karena merepresentasikan prediksi pada ukuran mesin nol, suatu nilai di luar domain data.', indent=True)
table_caption('Tabel 6. Koefisien model final M₃ (kubik)')
make_table([
    ['Suku', 'Koefisien', 'Std error', 't', 'p-value', 'CI 95%'],
    ['const', '4,6243', '0,1878', '24,62', '< 0,0001', '[4,256 ; 4,992]'],
    ['engine_size', '1,8027', '0,1612', '11,19', '< 0,0001', '[1,487 ; 2,119]'],
    ['engine_size²', '0,1044', '0,0425', '2,46', '0,0141', '[0,021 ; 0,188]'],
    ['engine_size³', '−0,0178', '0,0035', '−5,08', '< 0,0001', '[−0,025 ; −0,011]'],
], widths=[1.3, 0.9, 0.8, 0.6, 0.9, 1.5])

subheading('4.6 Validasi Asumsi OLS (Tahap 4)')
para('Anderson-Darling test menghasilkan statistik 149,99 yang jauh melampaui nilai kritis 5% sebesar '
     '0,787, sehingga normalitas residu ditolak secara formal. Breusch-Pagan test menghasilkan LM = '
     '1414,34 dengan p ≈ 2,3 × 10⁻³⁰⁶, sehingga homoskedastisitas juga ditolak. Namun, kedua hasil ini '
     'harus dibaca dengan hati-hati: pada ukuran sampel besar (n ≈ 18 ribu), kedua uji memiliki daya '
     'yang sangat tinggi sehingga mampu mendeteksi penyimpangan terkecil sekalipun. Penolakan formal '
     'tidak otomatis membatalkan model.', indent=True)
para('Gambar 4 menyajikan diagnostik visual yang lebih substantif. Pada Q-Q plot, bagian tengah residu '
     'mengikuti garis 45° dengan penyimpangan hanya pada ekor, pola yang umum dan dapat diterima untuk '
     'data rekayasa dengan skewness kanan. Pada plot residuals vs fitted, awan residu sedikit melebar '
     'seiring naiknya nilai prediksi, konsisten dengan heteroskedastisitas yang ditandai Breusch-Pagan. '
     'Implikasinya, estimasi titik koefisien tetap tak bias, tetapi bila diperlukan inferensi presisi '
     'pada standard error, solusi standarnya adalah refit dengan standard error robust terhadap '
     'heteroskedastisitas (HC3).', indent=True)
figure('figures/cell67_img7.png', 'Gambar 4. Diagnostik residu model final: Q-Q plot (kiri) dan plot '
       'residuals vs fitted (kanan).', width=6.0)

subheading('4.7 Evaluasi Generalisasi')
para('Tabel 7 membandingkan kinerja model final pada subset pelatihan dan uji. Selisih metrik pelatihan '
     'terhadap uji sangat kecil (R² berselisih sekitar 0,001 dan RMSE sekitar 0,03 L/100 km), sehingga '
     'tidak ada indikasi overfitting, sesuai harapan untuk model dengan sedikit parameter pada sampel '
     'besar. RMSE uji ≈ 1,73 L/100 km dan MAE uji ≈ 1,26 L/100 km berarti prediksi tipikal meleset '
     'sekitar 1,3–1,7 L/100 km dari nilai sebenarnya.', indent=True)
table_caption('Tabel 7. Evaluasi generalisasi model final')
make_table([
    ['Metrik', 'Pelatihan', 'Uji', 'Selisih (latih − uji)'],
    ['R²', '0,6570', '0,6582', '−0,0012'],
    ['RMSE (L/100 km)', '1,6984', '1,7261', '−0,0277'],
    ['MAE (L/100 km)', '1,2440', '1,2575', '−0,0135'],
], widths=[1.7, 1.0, 1.0, 1.6])
para('Secara keseluruhan, ukuran mesin sendirian menjelaskan sekitar 66% variansi konsumsi bahan bakar '
     'gabungan. Sisanya berasal dari faktor yang tidak dimodelkan, seperti tipe bahan bakar, bobot '
     'kendaraan, aerodinamika, dan jenis drivetrain. Untuk pemakaian terapan, prediksi sebaiknya '
     'diperlakukan sebagai estimasi kasar karena masih menyisakan ketidakpastian per kendaraan yang '
     'cukup besar.', indent=True)

# =====================================================================
# 5. KESIMPULAN DAN SARAN
# =====================================================================
heading(5, 'Kesimpulan dan Saran')
para('Penelitian ini memodelkan konsumsi bahan bakar gabungan kendaraan melalui kerangka keputusan '
     'bertingkat. Seleksi fitur Tahap 1 menyisakan satu prediktor, yaitu ukuran mesin, setelah jumlah '
     'silinder dibuang karena hampir kolinier (r ≈ 0,91) dan tahun model dibuang karena korelasinya '
     'lemah. Dua Nested F-Test secara berurutan menolak hipotesis nol sehingga model kubik terpilih '
     'sebagai model final. Model tersebut menjelaskan sekitar 66% variansi (R² uji ≈ 0,658) dengan RMSE '
     'uji ≈ 1,73 L/100 km dan tanpa indikasi overfitting. Bentuk kurva yang naik lalu melandai menegaskan '
     'pola diminishing returns pada hubungan ukuran mesin dengan konsumsi bahan bakar. Uji normalitas dan '
     'homoskedastisitas ditolak secara formal, namun hal ini merupakan konsekuensi wajar dari ukuran '
     'sampel besar dan tidak membatalkan estimasi titik koefisien.', indent=True)
para('Penelitian ini memiliki keterbatasan. Model bersifat prediktif dan asosiatif, bukan kausal; '
     'koefisien ukuran mesin menyerap segala sesuatu yang berkorelasi dengannya tetapi tidak ikut '
     'dimodelkan. Untuk penelitian lanjutan, disarankan menambahkan prediktor lain yang relevan (tipe '
     'bahan bakar, bobot, aerodinamika, drivetrain), menerapkan standard error robust HC3 untuk '
     'inferensi yang lebih presisi, serta mengeksplorasi model nonlinier alternatif seperti regresi '
     'spline atau generalized additive model guna menangkap struktur yang lebih kaya.', indent=True)

# =====================================================================
# UCAPAN TERIMA KASIH
# =====================================================================
heading(None, 'Ucapan Terima Kasih')
para('Penulis mengucapkan terima kasih kepada dosen pengampu mata kuliah Probabilitas dan Statistika '
     'atas bimbingan selama penyusunan penelitian ini, serta kepada penyedia dataset publik Fuel '
     'Consumption 2000–2022 yang memungkinkan analisis ini dilakukan.', indent=True)

# =====================================================================
# DAFTAR PUSTAKA
# =====================================================================
heading(None, 'Daftar Pustaka')
refs = [
    'Anderson, T. W., & Darling, D. A. (1952). Asymptotic theory of certain "goodness of fit" criteria '
    'based on stochastic processes. The Annals of Mathematical Statistics, 23(2), 193–212.',
    'Breusch, T. S., & Pagan, A. R. (1979). A simple test for heteroscedasticity and random coefficient '
    'variation. Econometrica, 47(5), 1287–1294.',
    'Draper, N. R., & Smith, H. (1998). Applied regression analysis (3rd ed.). New York: John Wiley & Sons.',
    'James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). An introduction to statistical learning '
    'with applications in R. New York: Springer.',
    'Montgomery, D. C., Peck, E. A., & Vining, G. G. (2012). Introduction to linear regression analysis '
    '(5th ed.). Hoboken, NJ: John Wiley & Sons.',
    'Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling with Python. '
    'Proceedings of the 9th Python in Science Conference, 92–96.',
    'Spearman, C. (1904). The proof and measurement of association between two things. The American '
    'Journal of Psychology, 15(1), 72–101.',
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(r); _set_run(run, 10)

# =====================================================================
# LAMPIRAN
# =====================================================================
heading(None, 'Lampiran')
para('Bagian ini memuat visualisasi berukuran besar yang mendukung analisis pada bagian utama.',
     indent=True)
figure('figures/cell24_img3.png', 'Gambar A1. Matriks korelasi Pearson antar variabel.', width=4.3)
figure('figures/cell27_img4.png', 'Gambar A2. Matriks korelasi Spearman antar variabel.', width=4.3)
figure('figures/cell38_img5.png', 'Gambar A3. Pair plot seluruh variabel pada sampel acak 3.000 baris. '
       'Panel ukuran mesin vs jumlah silinder memperlihatkan hubungan nyaris deterministik yang menjadi '
       'dasar pembuangan jumlah silinder.', width=5.0)

doc.save('paper_ALP_Stats.docx')
print('saved paper_ALP_Stats.docx')
