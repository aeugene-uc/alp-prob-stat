# -*- coding: utf-8 -*-
"""Builder paper dari AI.ipynb (regresi linear berganda + one-hot + HC3).
Gaya: baku, mudah dipahami, tanpa ';' dan tanpa '-' sebagai penyambung kalimat.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = normal.paragraph_format
pf.space_after = Pt(0)
pf.line_spacing = 1.15

for s in doc.sections:
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

def _set_run(r, size=11, bold=False, italic=False):
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    rpr = r._element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn('w:eastAsia'), 'Times New Roman')

def para(text='', align='justify', size=11, bold=False, italic=False,
         indent=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    a = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
         'left': WD_ALIGN_PARAGRAPH.LEFT, 'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.alignment = a
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    if text:
        r = p.add_run(text); _set_run(r, size, bold, italic)
    return p

def runs_para(segments, align='justify', indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    for t, b, i in segments:
        r = p.add_run(t); _set_run(r, 11, b, i)
    return p

def bullet(segments, space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    for t, b, i in segments:
        r = p.add_run(t); _set_run(r, 11, b, i)
    return p

def heading(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num}. {text}' if num else text)
    _set_run(r, 12, bold=True)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); _set_run(r, 11, bold=True, italic=True)
    return p

def _horizontal_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom', 'insideH'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        borders.append(el)
    for edge in ('left', 'right', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)

def caption(text, space_before=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); _set_run(r, 10, bold=True)

def make_table(rows, header=True, widths=None, align_center=True, fontsize=9):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _horizontal_borders(t)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            if align_center and ci > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align_center and ci == 0 and ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(val))
            _set_run(r, fontsize, bold=(header and ri == 0))
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(rows)):
                t.cell(ri, ci).width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, cap, width=5.8):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(); r.add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    rr = c.add_run(cap); _set_run(rr, 10, bold=False)

FIG = 'fig_ai/'

# =====================================================================
# TITLE BLOCK
# =====================================================================
para('Pemodelan Konsumsi Bahan Bakar Gabungan Kendaraan Menggunakan '
     'Regresi Linear Berganda dengan One-Hot Encoding',
     align='center', size=16, bold=True, space_after=4)
para('Modeling Combined Vehicle Fuel Consumption Using Multiple Linear '
     'Regression with One-Hot Encoding',
     align='center', size=12, bold=False, italic=True, space_after=8)

para('Rei Putra Soemanto, Amadeus Eugene Dirgantara, Jason Tio',
     align='center', size=11, bold=True, space_after=2)
para('Program Studi Informatika, Universitas Ciputra, Surabaya, Jawa Timur, Indonesia',
     align='center', size=10, space_after=1)
para('NIM: 0706012410060, 0706012410063, 0706012410006',
     align='center', size=10, space_after=1)
para('Tugas Mata Kuliah Probabilitas dan Statistika',
     align='center', size=10, italic=True, space_after=8)

# ----- Abstrak (Indonesia) -----
runs_para([
    ('Abstrak. ', True, False),
    ('Konsumsi bahan bakar kendaraan berkaitan langsung dengan biaya operasional, emisi, dan '
     'kebijakan keberlanjutan, sehingga kemampuan memprediksinya dari spesifikasi teknis kendaraan '
     'memiliki nilai praktis. Penelitian ini memodelkan konsumsi bahan bakar gabungan (comb_l_100km, '
     'dalam L/100 km) menggunakan dataset Fuel Consumption 2000 sampai 2022 yang berisi 22.555 baris '
     'setelah pembersihan. Pendekatan yang dipakai adalah regresi linear berganda dengan one-hot '
     'encoding pada variabel kategorikal jenis bahan bakar, kelas kendaraan, dan tipe transmisi, '
     'ditambah prediktor numerik ukuran mesin, jumlah silinder, dan jumlah gigi. Analisis disusun '
     'secara bertahap mulai dari eksplorasi data, seleksi fitur, pembentukan model, sampai validasi '
     'asumsi. Tahun model dibuang karena korelasinya dengan respons mendekati nol. Jumlah silinder '
     'tetap dipertahankan meskipun hampir kolinier dengan ukuran mesin karena uji ablasi menunjukkan '
     'bahwa ia masih menambah daya prediksi. Suku kuadratik ukuran mesin sempat diuji namun tidak '
     'menaikkan R-squared pada data uji, sehingga model linear yang lebih sederhana dipilih. Model '
     'akhir menjelaskan sekitar 84 persen variansi konsumsi bahan bakar (R-squared uji sekitar 0,84) '
     'dengan RMSE uji sekitar 1,18 L/100 km dan tanpa indikasi overfitting. Penggerak utama konsumsi '
     'adalah jenis bahan bakar, kelas kendaraan, dan ukuran mesin. Uji Breusch-Pagan dan Jarque-Bera '
     'menolak homoskedastisitas dan normalitas secara formal. Penolakan ini wajar pada ukuran sampel '
     'besar dan tidak membatalkan estimasi titik koefisien. Standard error robust HC3 dipakai agar '
     'inferensi tetap sahih terhadap heteroskedastisitas.', False, False)
], space_after=4)
runs_para([('Kata kunci. ', True, False),
           ('regresi linear berganda, one-hot encoding, seleksi fitur, konsumsi bahan bakar, '
            'multikolinieritas, heteroskedastisitas, standard error robust HC3.', False, False)],
          space_after=6)

runs_para([
    ('Abstract. ', True, True),
    ('Vehicle fuel consumption is directly related to operating cost, emissions, and sustainability '
     'policy, so the ability to predict it from technical specifications has practical value. This '
     'study models combined fuel consumption (comb_l_100km, in L/100 km) using the Fuel Consumption '
     '2000 to 2022 dataset containing 22,555 rows after cleaning. The approach is multiple linear '
     'regression with one-hot encoding on the categorical variables fuel type, vehicle class, and '
     'transmission type, together with the numeric predictors engine size, cylinder count, and number '
     'of gears. The analysis follows a staged workflow from data exploration, feature selection, model '
     'building, to assumption validation. Model year is dropped because its correlation with the '
     'response is near zero. Cylinder count is retained even though it is nearly collinear with engine '
     'size because an ablation test shows that it still adds predictive power. A quadratic engine-size '
     'term was tested but did not raise the test R-squared, so the simpler linear model was chosen. The '
     'final model explains about 84 percent of fuel-consumption variance (test R-squared about 0.84) '
     'with a test RMSE of about 1.18 L/100 km and no sign of overfitting. The main drivers are fuel '
     'type, vehicle class, and engine size. The Breusch-Pagan and Jarque-Bera tests reject '
     'homoscedasticity and normality formally. These rejections are expected at large sample sizes and '
     'do not invalidate the coefficient point estimates. HC3 robust standard errors are used so that '
     'inference remains valid under heteroscedasticity.', False, True)
], space_after=4)
runs_para([('Keywords. ', True, True),
           ('multiple linear regression, one-hot encoding, feature selection, fuel consumption, '
            'multicollinearity, heteroscedasticity, HC3 robust standard errors.', False, True)],
          space_after=8)

# =====================================================================
# 1. PENDAHULUAN
# =====================================================================
heading(1, 'Pendahuluan')
para('Konsumsi bahan bakar kendaraan merupakan indikator penting yang berhubungan langsung dengan '
     'biaya operasional, emisi gas rumah kaca, dan kebijakan keberlanjutan transportasi. Kemampuan '
     'memprediksi konsumsi bahan bakar dari karakteristik teknis kendaraan bermanfaat untuk '
     'perencanaan armada, perbandingan biaya bagi konsumen, deteksi anomali pada data spesifikasi, '
     'sampai pelaporan keberlanjutan. Oleh karena itu pemodelan statistik yang tepat terhadap '
     'hubungan antara karakteristik kendaraan dan konsumsi bahan bakar menjadi kebutuhan yang relevan.',
     indent=True)
para('Karakteristik kendaraan terdiri atas besaran numerik dan kategori. Besaran numerik mencakup '
     'ukuran mesin, jumlah silinder, dan jumlah gigi transmisi. Kategori mencakup jenis bahan bakar, '
     'kelas kendaraan, dan tipe transmisi. Variabel kategori tidak dapat langsung masuk ke dalam '
     'regresi karena nilainya berupa label. Variabel tersebut perlu diubah menjadi kolom indikator '
     'melalui one-hot encoding agar pengaruh tiap kategori dapat diestimasi secara terpisah terhadap '
     'kategori acuan.', indent=True)
para('Pemodelan ini menghadapi beberapa tantangan. Tantangan pertama adalah multikolinieritas. '
     'Ukuran mesin dan jumlah silinder berkorelasi sangat tinggi sehingga koefisiennya sulit '
     'ditafsirkan secara terpisah. Tantangan kedua adalah keberadaan prediktor yang hampir tidak '
     'relevan, misalnya tahun model, yang dapat memperberat model tanpa menambah daya prediksi. '
     'Tantangan ketiga adalah bentuk hubungan. Hubungan ukuran mesin dengan konsumsi bahan bakar '
     'mungkin sedikit melengkung sehingga perlu diperiksa apakah suku polinomial benar-benar '
     'diperlukan. Tantangan keempat adalah pemenuhan asumsi Ordinary Least Squares (OLS) yang menjadi '
     'dasar inferensi koefisien.', indent=True)
runs_para([
    ('Penelitian ini menjawab empat pertanyaan yang saling berkaitan. ', False, False),
    ('Pertama', True, False),
    (', fitur kendaraan mana yang sebaiknya dipertahankan sebagai prediktor. ', False, False),
    ('Kedua', True, False),
    (', bagaimana variabel kategori diolah agar informasinya terpakai tanpa menimbulkan kolinieritas '
     'sempurna. ', False, False),
    ('Ketiga', True, False),
    (', apakah model linear sudah memadai atau perlu suku polinomial untuk menangkap kelengkungan. ',
     False, False),
    ('Keempat', True, False),
    (', apakah model akhir memenuhi asumsi OLS dan bagaimana menangani pelanggaran yang muncul. Untuk '
     'menjawabnya, penelitian ini menerapkan kerangka analisis bertahap yang dimulai dari eksplorasi '
     'data, dilanjutkan seleksi fitur berbasis korelasi, pembentukan model dengan one-hot encoding, '
     'pemilihan antara model linear dan polinomial, validasi asumsi secara formal dan visual, serta '
     'penyempurnaan inferensi dengan standard error robust.', False, False)
], indent=True)

# =====================================================================
# 2. TINJAUAN PUSTAKA
# =====================================================================
heading(2, 'Tinjauan Pustaka')
subheading('2.1 Regresi Linear Berganda')
para('Regresi linear berganda memodelkan respons sebagai jumlah dari konstanta dan kontribusi linear '
     'tiap prediktor. Koefisien tiap prediktor menyatakan perubahan rata-rata respons untuk kenaikan '
     'satu satuan prediktor tersebut, dengan prediktor lain ditahan tetap. Model ini diestimasi '
     'menggunakan metode kuadrat terkecil yang meminimalkan jumlah kuadrat selisih antara nilai '
     'sebenarnya dan nilai prediksi. Kesederhanaan dan kemudahan interpretasinya membuat model ini '
     'menjadi titik awal yang wajar sebelum mempertimbangkan bentuk yang lebih kompleks.', indent=True)
subheading('2.2 One-Hot Encoding dan Kategori Acuan')
para('One-hot encoding mengubah satu variabel kategori menjadi sekumpulan kolom indikator bernilai '
     'nol atau satu. Untuk variabel dengan k kategori, dibuat k dikurang satu kolom indikator dan satu '
     'kategori ditetapkan sebagai acuan. Penetapan kategori acuan diperlukan untuk menghindari dummy '
     'variable trap, yaitu kolinieritas sempurna antara kolom indikator dan konstanta. Koefisien tiap '
     'kolom indikator kemudian dibaca sebagai selisih rata-rata respons kategori tersebut terhadap '
     'kategori acuan.', indent=True)
subheading('2.3 Korelasi Pearson dan Spearman')
para('Koefisien korelasi Pearson mengukur kekuatan hubungan linear antara dua variabel. Koefisien '
     'korelasi Spearman mengukur kekuatan hubungan monoton karena dihitung pada peringkat, sehingga '
     'lebih tahan terhadap pencilan dan tidak mengandaikan bentuk hubungan yang lurus. Membandingkan '
     'keduanya bersifat informatif. Apabila Spearman jauh lebih besar daripada Pearson, terdapat '
     'indikasi kelengkungan monoton yang membuat Pearson meremehkan kekuatan hubungan. Apabila '
     'keduanya berdekatan, hubungan monotonnya nyaris lurus sehingga Pearson sah dibaca sebagai '
     'kekuatan linear.', indent=True)
subheading('2.4 Multikolinieritas')
para('Multikolinieritas terjadi ketika dua atau lebih prediktor saling berkorelasi sangat tinggi. '
     'Akibatnya, koefisien menjadi sensitif terhadap perubahan kecil pada data dan standard error '
     'membesar sehingga koefisien sulit ditafsirkan secara terpisah. Akan tetapi, kolinieritas tidak '
     'selalu menurunkan daya prediksi. Selama pola korelasi antar prediktor tetap berlaku pada data '
     'baru, kedua prediktor masih dapat dipertahankan apabila terbukti menambah akurasi prediksi.',
     indent=True)
subheading('2.5 Validasi Asumsi OLS dan Standard Error Robust HC3')
para('Inferensi pada OLS bertumpu pada beberapa asumsi tentang error, terutama homoskedastisitas '
     'yang berarti variansi error konstan dan normalitas residual. Homoskedastisitas dapat diuji '
     'dengan Breusch-Pagan test dan normalitas dengan Jarque-Bera test, keduanya dilengkapi '
     'diagnostik visual berupa Q-Q plot dan plot residual terhadap nilai prediksi. Pada ukuran '
     'sampel besar, kedua uji formal sangat sensitif sehingga hampir selalu menolak walaupun '
     'penyimpangannya kecil. Ketika heteroskedastisitas terkonfirmasi, estimasi titik koefisien '
     'tetap tidak bias, namun standard error klasik menjadi keliru. Solusi standarnya adalah '
     'standard error robust terhadap heteroskedastisitas. Varian HC3 dikenal paling konservatif '
     'sehingga aman dipakai sebagai pilihan default.', indent=True)

# =====================================================================
# 3. METODE
# =====================================================================
heading(3, 'Metode')
subheading('3.1 Dataset dan Variabel')
para('Penelitian ini menggunakan dataset publik Fuel Consumption 2000 sampai 2022 yang memuat 22.556 '
     'baris dan 13 kolom spesifikasi kendaraan. Variabel respons adalah konsumsi bahan bakar gabungan '
     'comb_l_100km yang diambil dari kolom COMB (L/100 km). Prediktor numerik yang dipertimbangkan '
     'adalah ukuran mesin dalam liter, jumlah silinder, tahun model, dan jumlah gigi transmisi. '
     'Prediktor kategori yang dipertimbangkan adalah jenis bahan bakar, kelas kendaraan, dan tipe '
     'transmisi.', indent=True)
subheading('3.2 Persiapan Data')
para('Data mentah diperiksa untuk nilai yang hilang dan tidak ditemukan satu pun nilai kosong pada '
     'seluruh kolom, sehingga tidak diperlukan imputasi. Selanjutnya satu baris duplikat persis '
     'dihapus karena produsen kerap mencatat banyak konfigurasi trim yang identik sebagai baris '
     'tersendiri. Setelah penghapusan, tersisa 22.555 baris sebagai ukuran sampel efektif.', indent=True)
para('Label kelas kendaraan dirapikan karena memuat penulisan yang berbeda untuk makna yang sama, '
     'misalnya perbedaan kapitalisasi dan pemisah seperti "SUV - SMALL" dan "SUV: Small". Setelah '
     'penyeragaman, jumlah kelas menyusut dari 32 menjadi 17. Kolom transmisi dipecah menjadi dua '
     'fitur, yaitu tipe transmisi dan jumlah gigi. Transmisi variabel kontinu (CVT) tidak memiliki '
     'gigi tetap sehingga jumlah giginya ditetapkan nol sebagai konvensi.', indent=True)
subheading('3.3 Kerangka Analisis Bertahap')
para('Analisis dilakukan dalam lima tahap berurutan.', space_after=3)
bullet([('Tahap 1, eksplorasi data. ', True, False),
        ('Memetakan peran tiap kolom, meringkas distribusi, dan memeriksa pola hubungan melalui pair '
         'plot serta boxplot per kategori.', False, False)])
bullet([('Tahap 2, seleksi fitur. ', True, False),
        ('Menilai korelasi tiap prediktor terhadap respons dan korelasi antarprediktor untuk '
         'memutuskan fitur yang dipertahankan atau dibuang.', False, False)])
bullet([('Tahap 3, pembentukan model. ', True, False),
        ('Menerapkan one-hot encoding, membagi data menjadi subset latih dan uji, lalu membandingkan '
         'model linear dengan model polinomial.', False, False)])
bullet([('Tahap 4, validasi asumsi. ', True, False),
        ('Menguji homoskedastisitas dengan Breusch-Pagan dan normalitas dengan Jarque-Bera, dilengkapi '
         'diagnostik visual.', False, False)])
bullet([('Tahap 5, penyempurnaan inferensi. ', True, False),
        ('Mengganti standard error klasik dengan standard error robust HC3 dan melaporkan koefisien '
         'akhir beserta selang kepercayaannya.', False, False)])
subheading('3.4 Aturan Seleksi Fitur')
para('Prediktor dinilai berdasarkan kekuatan korelasinya dengan respons dan tingkat kolinieritasnya '
     'dengan prediktor lain. Prediktor dengan korelasi marginal yang dapat diabaikan terhadap respons '
     'menjadi kandidat untuk dibuang. Untuk prediktor yang saling kolinier, keputusan tidak hanya '
     'berdasar korelasi, melainkan juga diperiksa secara empiris melalui uji ablasi pada data uji. '
     'Uji ablasi membandingkan model lengkap dengan model tanpa prediktor tertentu untuk melihat '
     'apakah prediktor itu benar-benar menambah daya prediksi.', indent=True)
subheading('3.5 Estimasi dan Evaluasi')
para('Data dibagi menjadi subset latih dan subset uji dengan rasio 80 berbanding 20 menggunakan '
     'random state tetap demi reproduktibilitas. Subset latih berisi 18.044 baris dan subset uji '
     'berisi 4.511 baris. Subset latih dipakai untuk fitting model, sedangkan subset uji ditahan dan '
     'hanya dipakai untuk menilai generalisasi. Model diestimasi dengan OLS menggunakan pustaka '
     'Python, yaitu pandas, statsmodels, dan scikit-learn. Kualitas prediksi diukur dengan R-squared, '
     'RMSE, dan MAE, dengan RMSE dan MAE dinyatakan dalam satuan asli respons L/100 km.', indent=True)

# =====================================================================
# 4. HASIL DAN PEMBAHASAN
# =====================================================================
heading(4, 'Hasil dan Pembahasan')

subheading('4.1 Statistik Deskriptif')
para('Tabel 1 menyajikan statistik deskriptif variabel respons dan prediktor numerik. Ukuran mesin '
     'berkisar sekitar 0,8 sampai 8,4 liter dengan median 3,0 liter. Jumlah silinder terpusat pada '
     'konfigurasi 4, 6, dan 8. Tahun model tersebar cukup merata pada rentang 2000 sampai 2022. '
     'Respons comb_l_100km condong ke kanan dengan mayoritas kendaraan pada rentang 9 sampai 13 '
     'L/100 km dan ekor tipis yang melampaui 20 L/100 km.', indent=True)
caption('Tabel 1. Statistik deskriptif variabel numerik (n = 22.555)')
make_table([
    ['Variabel', 'Mean', 'Std', 'Min', 'Q1', 'Median', 'Q3', 'Max'],
    ['Ukuran mesin (L)', '3,357', '1,335', '0,8', '2,3', '3,0', '4,2', '8,4'],
    ['Jumlah silinder', '5,854', '1,820', '2', '4', '6', '8', '16'],
    ['Tahun model', '2011,6', '6,298', '2000', '2006', '2012', '2017', '2022'],
    ['Jumlah gigi', '5,635', '1,957', '0', '5', '6', '6', '10'],
    ['comb_l_100km', '11,034', '2,911', '3,6', '9,1', '10,6', '12,7', '26,1'],
], widths=[1.6, 0.7, 0.7, 0.6, 0.6, 0.75, 0.6, 0.6])
para('Komposisi kategori juga ditinjau. Pada jenis bahan bakar, Regular gasoline dan Premium gasoline '
     'mendominasi dengan masing-masing 11.822 dan 9.315 kendaraan, sedangkan Diesel dan Natural gas '
     'tergolong langka dengan 314 dan 33 kendaraan. Pada kelas kendaraan, kelas Compact, Mid Size, dan '
     'SUV paling sering muncul. Pada tipe transmisi, transmisi otomatis paling banyak, diikuti '
     'transmisi otomatis dengan select shift dan transmisi manual.', indent=True)

subheading('4.2 Eksplorasi Visual')
para('Pair plot variabel numerik pada Gambar A1 di Lampiran memperlihatkan pola menanjak yang jelas '
     'antara ukuran mesin, jumlah silinder, dan konsumsi bahan bakar. Jumlah silinder dan jumlah gigi '
     'bernilai diskret sehingga sebaran titiknya membentuk pola vertikal. Tahun model menyebar hampir '
     'datar terhadap konsumsi bahan bakar yang menandakan korelasi mendekati nol. Hubungan tiap '
     'prediktor numerik dengan respons tampak mendekati garis lurus yang mendukung pemakaian model '
     'linear.', indent=True)
para('Boxplot konsumsi bahan bakar menurut tiap variabel kategori pada Gambar A2 sampai A4 di '
     'Lampiran menegaskan bahwa ketiga faktor kategori benar-benar membedakan konsumsi. Jenis bahan '
     'bakar Diesel paling hemat, sedangkan Ethanol E85 dan Natural gas paling boros. Kelas kendaraan '
     'menanjak dari sedan kecil ke van dan pikap. Tipe transmisi memperlihatkan perbedaan yang lebih '
     'landai namun tetap sistematis. Karena setiap faktor menggeser median konsumsi antar kategori, '
     'ketiganya layak dimasukkan ke dalam model melalui one-hot encoding.', indent=True)

subheading('4.3 Analisis Korelasi dan Seleksi Fitur')
para('Tabel 2 merangkum korelasi Pearson dan Spearman tiap prediktor numerik terhadap respons dan '
     'korelasi antara ukuran mesin dan jumlah silinder. Matriks korelasi lengkap disajikan pada '
     'Gambar A5 dan A6 di Lampiran. Ukuran mesin dan jumlah silinder berkorelasi kuat dengan respons. '
     'Tahun model praktis tidak berkorelasi dengan respons karena nilainya berada di dalam rentang '
     'yang dianggap dapat diabaikan.', indent=True)
caption('Tabel 2. Korelasi prediktor numerik (terhadap comb_l_100km, dan antarprediktor)')
make_table([
    ['Pasangan', 'Pearson', 'Spearman', 'Putusan'],
    ['Ukuran mesin terhadap respons', '0,807', '0,845', 'Berarti'],
    ['Jumlah silinder terhadap respons', '0,772', '0,818', 'Berarti'],
    ['Tahun model terhadap respons', '-0,068', '-0,073', 'Dapat diabaikan'],
    ['Ukuran mesin terhadap jumlah silinder', '0,913', '0,936', 'Kolinier'],
], widths=[3.0, 0.9, 0.9, 1.3])
para('Putusan korelasi diambil dengan membandingkan selang kepercayaan 95 persen koefisien terhadap '
     'ambang praktis sebesar 0,1, bukan dengan p-value. Pada ukuran sampel sekitar 22.555, p-value '
     'hampir selalu menyatakan signifikan karena uji p-value hanya menanyakan apakah korelasi tepat '
     'nol, sehingga korelasi sekecil apa pun cukup untuk menolak nol. Dengan demikian p-value tidak '
     'dapat membedakan korelasi yang berarti dari yang dapat diabaikan, dan selang kepercayaanlah yang '
     'menjadi dasar putusan.', indent=True)
para('Perbandingan Pearson dan Spearman menunjukkan selisih yang kecil dan positif pada pasangan '
     'prediktor terhadap respons, yaitu sekitar 0,04 untuk ukuran mesin dan 0,05 untuk jumlah silinder. '
     'Selisih sekecil ini menandakan kelengkungan monoton yang sangat ringan, sehingga Pearson tetap '
     'sah dibaca sebagai kekuatan hubungan linear. Temuan ini menjadi dasar untuk tetap menguji suku '
     'kuadratik secara khusus, namun keputusannya diserahkan pada perbandingan pada data uji.', indent=True)
para('Tabel 3 mendokumentasikan keputusan seleksi fitur. Ukuran mesin dipertahankan karena paling '
     'kuat berhubungan dengan respons. Jumlah silinder dipertahankan meskipun hampir kolinier dengan '
     'ukuran mesin karena korelasinya bukan sempurna, dan uji ablasi pada Subbab 4.7 membuktikan ia '
     'masih menambah daya prediksi. Tahun model dibuang karena korelasinya dapat diabaikan. Jumlah '
     'gigi dipertahankan karena boxplot menunjukkan transmisi berkaitan dengan konsumsi. Ketiga '
     'variabel kategori dipertahankan untuk diolah dengan one-hot encoding.', indent=True)
caption('Tabel 3. Keputusan seleksi fitur')
make_table([
    ['Prediktor', 'Keputusan', 'Alasan ringkas'],
    ['Ukuran mesin', 'Pertahankan', 'Korelasi paling kuat dengan respons'],
    ['Jumlah silinder', 'Pertahankan', 'Kolinier, namun terbukti menambah prediksi (ablasi)'],
    ['Tahun model', 'Buang', 'Korelasi dengan respons dapat diabaikan'],
    ['Jumlah gigi', 'Pertahankan', 'Transmisi berkaitan dengan konsumsi'],
    ['Jenis bahan bakar', 'Pertahankan', 'Median konsumsi bergeser antar kategori'],
    ['Kelas kendaraan', 'Pertahankan', 'Median konsumsi bergeser antar kategori'],
    ['Tipe transmisi', 'Pertahankan', 'Median konsumsi bergeser antar kategori'],
], widths=[1.5, 1.2, 3.3], align_center=False)

subheading('4.4 Pembentukan Model dan Pemilihan Bentuk Fungsional')
para('Variabel kategori diubah menjadi kolom indikator melalui one-hot encoding dengan kategori acuan '
     'Diesel untuk bahan bakar, Compact untuk kelas kendaraan, dan transmisi otomatis untuk tipe '
     'transmisi. Proses ini menghasilkan 24 kolom indikator. Bersama tiga prediktor numerik, model '
     'linear memakai 27 fitur ditambah satu konstanta. Data kemudian dibagi menjadi subset latih '
     'berukuran 18.044 baris dan subset uji berukuran 4.511 baris.', indent=True)
para('Karena perbandingan Pearson dan Spearman menandai kelengkungan kecil pada ukuran mesin, model '
     'linear dibandingkan dengan model polinomial yang menambahkan suku kuadratik ukuran mesin. Tabel '
     '4 menunjukkan bahwa pada data uji kedua model praktis setara, padahal model polinomial memakai '
     'satu parameter lebih banyak. Karena tambahan suku kuadratik tidak menaikkan kemampuan '
     'generalisasi, model linear yang lebih sederhana dipilih sesuai prinsip parsimoni.', indent=True)
caption('Tabel 4. Perbandingan model linear dan polinomial')
make_table([
    ['Model', 'Jumlah parameter', 'R-squared latih', 'R-squared uji'],
    ['Linear', '28', '0,8342', '0,8365'],
    ['Polinomial (ukuran mesin kuadratik)', '29', '0,8344', '0,8364'],
], widths=[2.8, 1.3, 1.3, 1.2])

subheading('4.5 Model Final dan Interpretasi Koefisien')
para('Tabel 5 menyajikan koefisien model linear final beserta standard error, p-value, dan selang '
     'kepercayaan 95 persen. Seluruh nilai berbasis kovarians robust HC3 yang dijelaskan pada Subbab '
     '4.6. Koefisien numerik dibaca sebagai perubahan konsumsi untuk kenaikan satu satuan prediktor. '
     'Koefisien kategori dibaca sebagai selisih konsumsi terhadap kategori acuan. Konstanta mewakili '
     'kategori acuan saat semua prediktor numerik bernilai nol dan tidak ditafsirkan secara harfiah '
     'karena berada di luar rentang data.', indent=True)
caption('Tabel 5. Koefisien model linear final dengan standard error robust HC3')
coef_rows = [
    ['Variabel', 'Koefisien', 'SE (HC3)', 'p-value', 'CI bawah', 'CI atas'],
    ['Konstanta', '3,1407', '0,0868', '< 0,001', '2,971', '3,311'],
    ['Ukuran mesin', '0,6438', '0,0264', '< 0,001', '0,592', '0,696'],
    ['Jumlah silinder', '0,5740', '0,0185', '< 0,001', '0,538', '0,610'],
    ['Jumlah gigi', '-0,0518', '0,0075', '< 0,001', '-0,066', '-0,037'],
    ['Bahan bakar: Ethanol E85 (E)', '5,4887', '0,0758', '< 0,001', '5,340', '5,637'],
    ['Bahan bakar: Natural gas (N)', '3,7854', '0,2542', '< 0,001', '3,287', '4,284'],
    ['Bahan bakar: Premium gasoline (Z)', '2,0951', '0,0574', '< 0,001', '1,983', '2,208'],
    ['Bahan bakar: Regular gasoline (X)', '1,4221', '0,0569', '< 0,001', '1,311', '1,534'],
    ['Kelas: Full size', '0,0155', '0,0435', '0,721', '-0,070', '0,101'],
    ['Kelas: Mid size', '0,0016', '0,0304', '0,958', '-0,058', '0,061'],
    ['Kelas: Minicompact', '-0,1952', '0,0441', '< 0,001', '-0,282', '-0,109'],
    ['Kelas: Minivan', '1,0503', '0,0490', '< 0,001', '0,954', '1,147'],
    ['Kelas: Pickup truck small', '2,0659', '0,0553', '< 0,001', '1,958', '2,174'],
    ['Kelas: Pickup truck standard', '1,9945', '0,0431', '< 0,001', '1,910', '2,079'],
    ['Kelas: Special purpose vehicle', '1,9189', '0,0911', '< 0,001', '1,740', '2,098'],
    ['Kelas: Station wagon mid size', '0,5699', '0,0469', '< 0,001', '0,478', '0,662'],
    ['Kelas: Station wagon small', '0,3136', '0,0406', '< 0,001', '0,234', '0,393'],
    ['Kelas: Subcompact', '0,0939', '0,0359', '0,009', '0,024', '0,164'],
    ['Kelas: SUV', '1,5188', '0,0365', '< 0,001', '1,447', '1,590'],
    ['Kelas: SUV small', '1,3675', '0,0364', '< 0,001', '1,296', '1,439'],
    ['Kelas: SUV standard', '1,7452', '0,0527', '< 0,001', '1,642', '1,848'],
    ['Kelas: Two seater', '0,4444', '0,0598', '< 0,001', '0,327', '0,562'],
    ['Kelas: Van cargo', '2,6157', '0,0718', '< 0,001', '2,475', '2,757'],
    ['Kelas: Van passenger', '3,9195', '0,1053', '< 0,001', '3,713', '4,126'],
    ['Transmisi: Automated manual (AM)', '0,4672', '0,0573', '< 0,001', '0,355', '0,580'],
    ['Transmisi: Automatic select shift (AS)', '0,1310', '0,0224', '< 0,001', '0,087', '0,175'],
    ['Transmisi: Continuously variable (AV)', '-1,6222', '0,0646', '< 0,001', '-1,749', '-1,496'],
    ['Transmisi: Manual (M)', '0,1623', '0,0253', '< 0,001', '0,113', '0,212'],
]
make_table(coef_rows, widths=[2.7, 0.85, 0.75, 0.75, 0.7, 0.7], align_center=False, fontsize=8.5)
para('Pembacaan koefisien menghasilkan beberapa temuan yang masuk akal secara teknis. Ukuran mesin '
     'dan jumlah silinder sama-sama positif, sehingga mesin yang lebih besar dan bersilinder lebih '
     'banyak cenderung lebih boros. Karena keduanya sangat berkaitan, koefisiennya dibaca bersama '
     'sebagai ukuran kapasitas mesin. Jumlah gigi bertanda negatif yang menandakan transmisi dengan '
     'lebih banyak rasio gigi cenderung sedikit lebih hemat. Pada jenis bahan bakar, Diesel paling '
     'hemat sebagai acuan, sedangkan Ethanol E85 paling boros, diikuti Natural gas, Premium, dan '
     'Regular. Pada tipe transmisi, transmisi variabel kontinu paling hemat. Pada kelas kendaraan, '
     'terdapat gradien yang jelas dari mobil kecil sampai kendaraan besar seperti van dan pikap.',
     indent=True)

subheading('4.6 Validasi Asumsi dan Penyempurnaan Inferensi HC3')
para('Validasi asumsi dilakukan pada homoskedastisitas dan normalitas. Asumsi independensi tidak '
     'diuji secara formal karena data bersifat cross-sectional sehingga urutan baris tidak bermakna. '
     'Tabel 6 merangkum hasil uji formal beserta ukuran magnitudo penyimpangannya. Breusch-Pagan '
     'menolak homoskedastisitas dan Jarque-Bera menolak normalitas. Kedua penolakan ini wajar pada '
     'ukuran sampel besar karena statistik uji membesar seiring jumlah data, sehingga penyimpangan '
     'sekecil apa pun ikut ditolak. Oleh karena itu pembacaan dilengkapi dengan magnitudo, yaitu '
     'skewness, excess kurtosis, dan proporsi variansi residual kuadrat yang dijelaskan prediktor.',
     indent=True)
caption('Tabel 6. Hasil uji asumsi dan magnitudo penyimpangan')
make_table([
    ['Aspek', 'Uji dan magnitudo', 'Pembacaan'],
    ['Homoskedastisitas', 'Breusch-Pagan, R-squared auxiliary = 0,083', 'Heteroskedastik ringan sampai sedang'],
    ['Normalitas (kemiringan)', 'Skewness = 0,42', 'Kemiringan dapat diabaikan'],
    ['Normalitas (ekor)', 'Excess kurtosis = 2,14', 'Ekor lebih tebal dari normal'],
], widths=[1.7, 2.7, 2.0], align_center=False)
para('Diagnostik visual pada Gambar 1 memperkuat pembacaan tersebut. Pada Q-Q plot, bagian tengah '
     'residual mengikuti garis 45 derajat dengan penyimpangan terkumpul di ekor, pola yang umum untuk '
     'respons yang sedikit condong. Pada plot residual terhadap nilai prediksi, sebaran residual '
     'sedikit melebar seperti corong seiring naiknya nilai prediksi yang menandakan heteroskedastisitas. '
     'Tidak ada lengkungan sistematis, sehingga spesifikasi linear tampak memadai.', indent=True)
figure(FIG + 'diagnostic.png',
       'Gambar 1. Diagnostik residual model final. Q-Q plot (kiri) dan plot residual terhadap nilai '
       'prediksi (kanan).', width=6.0)
para('Karena heteroskedastisitas terkonfirmasi, standard error klasik diganti dengan standard error '
     'robust HC3. Penggantian ini tidak mengubah estimasi titik koefisien sama sekali dan hanya '
     'memperbaiki ukuran ketidakpastiannya. Tabel 7 memperlihatkan bahwa pada sebagian koefisien '
     'standard error berubah cukup nyata. Standard error ukuran mesin dan jumlah silinder naik sekitar '
     '40 persen, sedangkan sebagian koefisien bahan bakar justru turun. Perubahan ini menjadi bukti '
     'langsung bahwa standard error klasik kurang tepat saat ada heteroskedastisitas, sehingga HC3 '
     'merupakan koreksi yang sesuai. Varian HC3 dipilih karena paling konservatif, sedangkan pada '
     'ukuran sampel besar seperti data ini selisih antar varian HC hampir nol sehingga memilih yang '
     'paling aman tidak merugikan.', indent=True)
caption('Tabel 7. Perbandingan standard error klasik dan HC3 (sampel koefisien)')
make_table([
    ['Koefisien', 'SE klasik', 'SE HC3', 'Rasio HC3 terhadap klasik'],
    ['Ukuran mesin', '0,0185', '0,0264', '1,43'],
    ['Jumlah silinder', '0,0131', '0,0185', '1,41'],
    ['Bahan bakar: Premium gasoline (Z)', '0,0782', '0,0574', '0,73'],
    ['Kelas: Van passenger', '0,0859', '0,1053', '1,23'],
], widths=[2.6, 1.0, 1.0, 1.7], align_center=False)

subheading('4.7 Uji Ablasi dan Evaluasi Generalisasi')
para('Uji ablasi menegaskan dua keputusan seleksi fitur. Membuang jumlah silinder menurunkan '
     'R-squared uji sekitar 0,02, dari 0,8365 menjadi 0,8168. Penurunan yang konsisten pada data uji '
     'menandakan informasi itu nyata, bukan hasil overfitting. Sebaliknya, menambahkan kembali tahun '
     'model praktis tidak mengubah R-squared uji dan koefisiennya tidak signifikan dengan p-value '
     'sekitar 0,852. Hasil ini menguatkan keputusan mempertahankan jumlah silinder dan membuang tahun '
     'model dengan standar held-out yang sama.', indent=True)
para('Tabel 8 membandingkan kinerja model final pada subset latih dan subset uji. Selisih metrik '
     'antara latih dan uji sangat kecil, sehingga tidak ada indikasi overfitting. RMSE uji sekitar '
     '1,18 L/100 km dan MAE uji sekitar 0,89 L/100 km berarti prediksi tipikal meleset sekitar satu '
     'liter per 100 km dari nilai sebenarnya.', indent=True)
caption('Tabel 8. Evaluasi generalisasi model final')
make_table([
    ['Metrik', 'Latih', 'Uji', 'Selisih (latih dikurang uji)'],
    ['R-squared', '0,8342', '0,8365', '-0,0023'],
    ['RMSE (L/100 km)', '1,1857', '1,1757', '0,0101'],
    ['MAE (L/100 km)', '0,8962', '0,8941', '0,0020'],
], widths=[1.8, 1.0, 1.0, 1.9])
para('Secara keseluruhan, model menjelaskan sekitar 84 persen variansi konsumsi bahan bakar gabungan. '
     'Sisanya berasal dari faktor yang tidak dimodelkan, seperti bobot kendaraan, aerodinamika, dan '
     'jenis penggerak. Untuk pemakaian terapan, prediksi sebaiknya diperlakukan sebagai estimasi yang '
     'baik namun tetap menyisakan ketidakpastian per kendaraan.', indent=True)

# =====================================================================
# 5. KESIMPULAN DAN SARAN
# =====================================================================
heading(5, 'Kesimpulan dan Saran')
para('Penelitian ini memodelkan konsumsi bahan bakar gabungan kendaraan menggunakan regresi linear '
     'berganda dengan one-hot encoding melalui kerangka analisis bertahap. Seleksi fitur '
     'mempertahankan ukuran mesin, jumlah silinder, dan jumlah gigi sebagai prediktor numerik serta '
     'jenis bahan bakar, kelas kendaraan, dan tipe transmisi sebagai prediktor kategori. Tahun model '
     'dibuang karena korelasinya dapat diabaikan dan tidak menambah daya prediksi. Jumlah silinder '
     'dipertahankan karena uji ablasi membuktikan ia masih menyumbang akurasi meskipun kolinier '
     'dengan ukuran mesin. Suku kuadratik ukuran mesin tidak diperlukan karena tidak menaikkan '
     'R-squared uji, sehingga model linear yang dipilih. Model akhir menjelaskan sekitar 84 persen '
     'variansi dengan RMSE uji sekitar 1,18 L/100 km dan tanpa indikasi overfitting. Penggerak utama '
     'konsumsi adalah jenis bahan bakar, kelas kendaraan, dan ukuran mesin.', indent=True)
para('Penelitian ini memiliki keterbatasan. Uji Breusch-Pagan dan Jarque-Bera menolak '
     'homoskedastisitas dan normalitas secara formal. Penolakan ini wajar pada ukuran sampel besar '
     'dan tidak membatalkan estimasi titik koefisien. Standard error robust HC3 sudah menangani '
     'heteroskedastisitas, sehingga p-value dan selang kepercayaan yang dilaporkan dapat ditafsirkan '
     'sebagaimana mestinya. Non-normalitas residual yang didorong oleh ekor tebal tidak fatal bagi '
     'prediksi titik berkat Central Limit Theorem, namun membuat interval prediksi kurang terkalibrasi '
     'di bagian ekor. Asumsi independensi tidak diuji, sehingga kemungkinan korelasi dalam grup antar '
     'varian dari model kendaraan yang sama masih terbuka. Model ini bersifat asosiatif, bukan kausal.',
     indent=True)
para('Untuk penelitian lanjutan disarankan beberapa hal. Pertama, menambahkan prediktor lain yang '
     'relevan seperti bobot kendaraan, aerodinamika, dan jenis penggerak untuk menaikkan daya '
     'prediksi. Kedua, menerapkan standard error berbasis cluster per model kendaraan untuk menangani '
     'kemungkinan korelasi dalam grup. Ketiga, mengeksplorasi metode yang memberikan interval prediksi '
     'lebih terkalibrasi pada data dengan ekor tebal.', indent=True)

# =====================================================================
# UCAPAN TERIMA KASIH
# =====================================================================
heading(None, 'Ucapan Terima Kasih')
para('Penulis mengucapkan terima kasih kepada dosen pengampu mata kuliah Probabilitas dan Statistika '
     'atas bimbingan selama penyusunan penelitian ini, serta kepada penyedia dataset publik Fuel '
     'Consumption 2000 sampai 2022 yang memungkinkan analisis ini dilakukan.', indent=True)

# =====================================================================
# DAFTAR PUSTAKA
# =====================================================================
heading(None, 'Daftar Pustaka')
refs = [
    'Breusch, T. S., & Pagan, A. R. (1979). A simple test for heteroscedasticity and random '
    'coefficient variation. Econometrica, 47(5), 1287 sampai 1294.',
    'Jarque, C. M., & Bera, A. K. (1980). Efficient tests for normality, homoscedasticity and serial '
    'independence of regression residuals. Economics Letters, 6(3), 255 sampai 259.',
    'James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). An introduction to statistical '
    'learning with applications in R. New York: Springer.',
    'MacKinnon, J. G., & White, H. (1985). Some heteroskedasticity-consistent covariance matrix '
    'estimators with improved finite sample properties. Journal of Econometrics, 29(3), 305 sampai 325.',
    'Montgomery, D. C., Peck, E. A., & Vining, G. G. (2012). Introduction to linear regression '
    'analysis (5th ed.). Hoboken, NJ: John Wiley & Sons.',
    'Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling with '
    'Python. Proceedings of the 9th Python in Science Conference, 92 sampai 96.',
    'Spearman, C. (1904). The proof and measurement of association between two things. The American '
    'Journal of Psychology, 15(1), 72 sampai 101.',
]
for rtext in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(rtext); _set_run(run, 10)

# =====================================================================
# LAMPIRAN
# =====================================================================
doc.add_page_break()
heading(None, 'Lampiran')
para('Bagian ini memuat visualisasi pendukung analisis pada bagian utama.', indent=True, space_after=8)
figure(FIG + 'pairplot.png',
       'Gambar A1. Pair plot variabel numerik pada sampel acak 3.000 baris.', width=5.4)
figure(FIG + 'box_fuel.png',
       'Gambar A2. Distribusi comb_l_100km menurut jenis bahan bakar.', width=5.6)
figure(FIG + 'box_vehicle_class.png',
       'Gambar A3. Distribusi comb_l_100km menurut kelas kendaraan.', width=5.6)
figure(FIG + 'box_transmission.png',
       'Gambar A4. Distribusi comb_l_100km menurut tipe transmisi.', width=5.6)
figure(FIG + 'pearson.png',
       'Gambar A5. Matriks korelasi Pearson antar variabel numerik.', width=4.6)
figure(FIG + 'spearman.png',
       'Gambar A6. Matriks korelasi Spearman antar variabel numerik.', width=4.6)

OUT = r'ALP Statistics & Probability.docx'
doc.save(OUT)
print('saved', OUT)
